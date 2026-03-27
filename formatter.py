"""
JGTTR Word → LaTeX Formatter
Converts free-format Word documents to JGTTR journal LaTeX template.
Also supports structured form-based input (generate_latex_from_form).
"""

import re
import os
import io
import zipfile
from docx import Document
from docx.oxml.ns import qn


# ── LaTeX special-character escaping ──────────────────────────────────────────
_LATEX_ESCAPE = str.maketrans({
    '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
    '_': r'\_', '{': r'\{', '}': r'\}',
    '~': r'\textasciitilde{}', '^': r'\textasciicircum{}',
    '\\': r'\textbackslash{}',
})

def escape(text: str) -> str:
    if not text:
        return ''
    return text.translate(_LATEX_ESCAPE)


# ── Heading / section detection ───────────────────────────────────────────────
SECTION_MAP = {
    # TR keys
    'giriş': 'Giriş / Introduction',
    'introduction': 'Giriş / Introduction',
    'literatür': 'Literatür Taraması / Literature Review',
    'literature': 'Literatür Taraması / Literature Review',
    'literature review': 'Literatür Taraması / Literature Review',
    'kuramsal çerçeve': 'Literatür Taraması / Literature Review',
    'yöntem': 'Yöntem / Methodology',
    'yöntem ve teknik': 'Yöntem / Methodology',
    'methodology': 'Yöntem / Methodology',
    'method': 'Yöntem / Methodology',
    'materials and methods': 'Yöntem / Methodology',
    'bulgular': 'Bulgular / Findings',
    'findings': 'Bulgular / Findings',
    'results': 'Bulgular / Findings',
    'tartışma': 'Tartışma / Discussion',
    'discussion': 'Tartışma / Discussion',
    'sonuç': 'Sonuç / Conclusion',
    'conclusion': 'Sonuç / Conclusion',
    'sonuç ve öneriler': 'Sonuç / Conclusion',
    'conclusions': 'Sonuç / Conclusion',
    'kaynakça': '__REFERENCES__',
    'kaynaklar': '__REFERENCES__',
    'references': '__REFERENCES__',
    'bibliography': '__REFERENCES__',
}

ABSTRACT_KEYS = {'abstract', 'özet', 'öz'}
KEYWORD_KEYS  = {'keywords', 'anahtar kelimeler', 'anahtar sözcükler', 'key words'}


def _para_text(para) -> str:
    return para.text.strip()


def _is_heading(para) -> bool:
    return para.style.name.startswith('Heading') or para.style.name.startswith('Başlık')


def _heading_level(para) -> int:
    name = para.style.name
    for part in name.split():
        if part.isdigit():
            return int(part)
    return 1


def _has_bold(para) -> bool:
    return any(run.bold for run in para.runs if run.text.strip())


def _para_to_latex(para) -> str:
    """Convert a paragraph with inline formatting to LaTeX."""
    parts = []
    for run in para.runs:
        t = escape(run.text)
        if not t:
            continue
        if run.bold and run.italic:
            t = r'\textbf{\textit{' + t + '}}'
        elif run.bold:
            t = r'\textbf{' + t + '}'
        elif run.italic:
            t = r'\textit{' + t + '}'
        parts.append(t)
    return ''.join(parts)


def _list_to_latex(paras, numbered: bool) -> str:
    env = 'enumerate' if numbered else 'itemize'
    opt = r'[leftmargin=1.2cm, label=\arabic*.]' if numbered else ''
    lines = [r'\begin{' + env + '}' + opt]
    for p in paras:
        lines.append(r'  \item ' + _para_to_latex(p))
    lines.append(r'\end{' + env + '}')
    return '\n'.join(lines)


# ── Word document parser ───────────────────────────────────────────────────────
def extract_from_docx(file_bytes: bytes) -> dict:
    """
    Returns a dict with keys:
      tr_title, en_title, tr_abstract, en_abstract,
      tr_keywords, en_keywords, sections (list of {title, level, latex}),
      references (list of str), raw_paragraphs
    """
    doc = Document(io.BytesIO(file_bytes))
    result = {
        'tr_title': '', 'en_title': '',
        'tr_abstract': '', 'en_abstract': '',
        'tr_keywords': '', 'en_keywords': '',
        'sections': [],
        'references': [],
    }

    paras = [p for p in doc.paragraphs]
    i = 0
    n = len(paras)

    current_section_title = None
    current_section_level = 1
    current_section_lines = []
    in_abstract = False
    in_keywords = False
    abstract_lang = None   # 'tr' or 'en'
    in_references = False

    def flush_section():
        if current_section_title is None:
            return
        result['sections'].append({
            'title': current_section_title,
            'level': current_section_level,
            'latex': '\n\n'.join(current_section_lines),
        })

    # ── first pass: find titles (usually first 1-3 paragraphs before abstract) ──
    first_title_found = False
    for p in paras[:8]:
        txt = _para_text(p)
        if not txt:
            continue
        key = txt.lower().strip().rstrip(':').strip()
        if key in ABSTRACT_KEYS or key in KEYWORD_KEYS:
            break
        if _is_heading(p) or len(txt) > 10:
            if not first_title_found:
                # Heuristic: if line looks like a Turkish title (no ASCII section name)
                if not any(key in txt.lower() for key in SECTION_MAP):
                    result['tr_title'] = escape(txt)
                    first_title_found = True
                    continue
            elif not result['en_title']:
                if not any(key in txt.lower() for key in SECTION_MAP):
                    result['en_title'] = escape(txt)
                    break

    # ── main pass ──
    i = 0
    while i < n:
        para = paras[i]
        txt = _para_text(para)
        key = txt.lower().strip().rstrip(':').strip()
        style = para.style.name

        # Skip empty
        if not txt:
            i += 1
            continue

        # ── References section ──
        if in_references:
            if txt:
                result['references'].append(escape(txt))
            i += 1
            continue

        # ── Heading detection ──
        if _is_heading(para) or (len(txt) < 80 and txt.isupper() and len(txt) > 3):
            mapped = SECTION_MAP.get(key)
            if mapped == '__REFERENCES__':
                flush_section()
                current_section_title = None
                current_section_lines = []
                in_references = True
                i += 1
                continue
            if mapped:
                flush_section()
                current_section_title = mapped
                current_section_level = _heading_level(para) if _is_heading(para) else 1
                current_section_lines = []
                in_abstract = False
                i += 1
                continue
            # Sub-heading inside a known section
            if current_section_title:
                flush_section()
                current_section_title = escape(txt)
                current_section_level = _heading_level(para) if _is_heading(para) else 2
                current_section_lines = []
                i += 1
                continue

        # ── Abstract heading ──
        if key in ABSTRACT_KEYS:
            in_abstract = True
            in_keywords = False
            # determine language from context
            if 'en' in key or key == 'abstract':
                abstract_lang = 'en'
            else:
                abstract_lang = 'tr'
            i += 1
            continue

        # ── Keywords line ──
        if key in KEYWORD_KEYS or txt.lower().startswith('keyword') or txt.lower().startswith('anahtar'):
            in_abstract = False
            in_keywords = True
            # might be on same line: "Keywords: foo, bar"
            colon_pos = txt.find(':')
            if colon_pos != -1:
                kw_val = txt[colon_pos+1:].strip()
                if abstract_lang == 'en':
                    result['en_keywords'] = escape(kw_val)
                else:
                    result['tr_keywords'] = escape(kw_val)
                in_keywords = False
            i += 1
            continue

        if in_keywords:
            if abstract_lang == 'en':
                result['en_keywords'] = escape(txt)
            else:
                result['tr_keywords'] = escape(txt)
            in_keywords = False
            i += 1
            continue

        # ── Abstract body ──
        if in_abstract:
            if abstract_lang == 'en':
                result['en_abstract'] += (' ' if result['en_abstract'] else '') + txt
            else:
                result['tr_abstract'] += (' ' if result['tr_abstract'] else '') + txt
            i += 1
            continue

        # ── Regular paragraph / list ──
        if current_section_title is None:
            i += 1
            continue

        # List paragraph
        if style.startswith('List') or para.style.name in ('List Paragraph', 'Liste Paragrafı'):
            # collect consecutive list items
            list_paras = [para]
            numbered = 'Number' in style or 'Numara' in style
            j = i + 1
            while j < n:
                np2 = paras[j]
                s2 = np2.style.name
                if s2.startswith('List') or s2 in ('List Paragraph',):
                    list_paras.append(np2)
                    j += 1
                else:
                    break
            current_section_lines.append(_list_to_latex(list_paras, numbered))
            i = j
            continue

        # Normal paragraph
        latex_line = _para_to_latex(para)
        if latex_line:
            current_section_lines.append(latex_line)
        i += 1

    flush_section()
    return result


# ── Author info parser ─────────────────────────────────────────────────────────
def parse_author_info(file_bytes: bytes, filename: str) -> list:
    """
    Parse author info from a .docx or .txt file.
    Expected format (one author per line/paragraph):
      Ad Soyad | Kurum | ORCID | email | sorumlu(evet/hayır)
    or Word table with columns: Ad Soyad, Kurum, ORCID, E-posta, Sorumlu
    Returns list of dicts.
    """
    authors = []

    if filename.lower().endswith('.docx'):
        doc = Document(io.BytesIO(file_bytes))
        # Try table first
        if doc.tables:
            tbl = doc.tables[0]
            for row in tbl.rows[1:]:  # skip header
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4 and cells[0]:
                    authors.append({
                        'name':        cells[0],
                        'affiliation': cells[1] if len(cells) > 1 else '',
                        'orcid':       cells[2] if len(cells) > 2 else '',
                        'email':       cells[3] if len(cells) > 3 else '',
                        'corresponding': len(cells) > 4 and cells[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                    })
        else:
            for para in doc.paragraphs:
                txt = para.text.strip()
                if not txt or txt.startswith('#'):
                    continue
                parts = [p.strip() for p in re.split(r'[|\t;]', txt)]
                if len(parts) >= 2:
                    authors.append({
                        'name':        parts[0],
                        'affiliation': parts[1] if len(parts) > 1 else '',
                        'orcid':       parts[2] if len(parts) > 2 else '',
                        'email':       parts[3] if len(parts) > 3 else '',
                        'corresponding': len(parts) > 4 and parts[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                    })
    else:
        # Plain text
        text = file_bytes.decode('utf-8', errors='ignore')
        for line in text.splitlines():
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            parts = [p.strip() for p in re.split(r'[|\t;]', line)]
            if len(parts) >= 2:
                authors.append({
                    'name':        parts[0],
                    'affiliation': parts[1] if len(parts) > 1 else '',
                    'orcid':       parts[2] if len(parts) > 2 else '',
                    'email':       parts[3] if len(parts) > 3 else '',
                    'corresponding': len(parts) > 4 and parts[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                })

    return authors


# ── LaTeX generation ───────────────────────────────────────────────────────────
SECTION_CMD = {1: r'\section', 2: r'\subsection', 3: r'\subsubsection'}

def _format_author_block(authors: list) -> str:
    parts = []
    for idx, a in enumerate(authors, 1):
        sup = str(idx)
        if a.get('corresponding'):
            sup += ',*'
        orcid = a.get('orcid', '')
        orcid_part = r'\,\orcidlink{' + orcid + '}' if orcid else ''
        parts.append(escape(a['name']) + r'\textsuperscript{' + sup + r'}' + orcid_part)
    return ',\n  '.join(parts)


def _format_affiliations(authors: list) -> str:
    lines = []
    for idx, a in enumerate(authors, 1):
        orcid = a.get('orcid', '')
        email = a.get('email', '')
        orcid_part = r' ORCID: \href{https://orcid.org/' + orcid + r'}{\mbox{' + orcid + r'}}.' if orcid else ''
        email_part = r' E-posta: \href{mailto:' + email + r'}{\mbox{' + email + r'}}' if email else ''
        lines.append(
            r'\textsuperscript{' + str(idx) + r'}' +
            escape(a.get('affiliation', '')) + '.' +
            orcid_part + email_part
        )
    return r'\\' + '\n  '.join(lines)


def _format_corresponding(authors: list) -> str:
    for a in authors:
        if a.get('corresponding'):
            name  = escape(a.get('name', ''))
            email = a.get('email', '')
            aff   = escape(a.get('affiliation', ''))
            ep    = r'. E-posta: \href{mailto:' + email + r'}{' + email + r'}' if email else ''
            return name + (', ' + aff if aff else '') + ep
    if authors:
        a = authors[0]
        name  = escape(a.get('name', ''))
        email = a.get('email', '')
        ep    = r'. E-posta: \href{mailto:' + email + r'}{' + email + r'}' if email else ''
        return name + ep
    return 'Yazar Adı, Kurum, E-posta'


def generate_latex(content: dict, authors: list, meta: dict) -> str:
    """Generate complete JGTTR-formatted LaTeX source."""

    tr_title   = content.get('tr_title') or r'Makalenin Türkçe Adı'
    en_title   = content.get('en_title') or r'Article Title in English'
    tr_abs     = content.get('tr_abstract') or r'Türkçe özet buraya yazılmalıdır.'
    en_abs     = content.get('en_abstract') or r'English abstract goes here.'
    tr_kw      = content.get('tr_keywords') or r'anahtar kelime 1; anahtar kelime 2'
    en_kw      = content.get('en_keywords') or r'keyword 1; keyword 2'

    year       = escape(meta.get('year', '2026'))
    volume     = escape(meta.get('volume', 'x'))
    issue      = escape(meta.get('issue', 'x'))
    start_page = escape(meta.get('start_page', 'xxx'))
    end_page   = escape(meta.get('end_page', 'xxx'))
    doi        = meta.get('doi', '')

    author_short = meta.get('author_short', '')
    if not author_short and authors:
        names = [a['name'].split()[-1] for a in authors]
        if len(names) == 1:
            author_short = escape(names[0])
        elif len(names) == 2:
            author_short = escape(names[0]) + r' \& ' + escape(names[1])
        else:
            author_short = escape(names[0]) + r' et al.'

    head_title = meta.get('head_title', '') or tr_title[:60]

    author_block  = _format_author_block(authors) if authors else r'Author One\textsuperscript{1,*}'
    affiliations  = _format_affiliations(authors) if authors else r'\textsuperscript{1}Kurum Adı, Şehir.'
    corresponding = _format_corresponding(authors) if authors else r'Yazar Adı, Kurum.'

    # ── body sections ──
    body_lines = []
    for sec in content.get('sections', []):
        title = sec['title']
        level = sec.get('level', 1)
        cmd   = SECTION_CMD.get(level, r'\section')
        body_lines.append(cmd + '{' + title + '}')
        if sec['latex']:
            body_lines.append('')
            body_lines.append(sec['latex'])
        body_lines.append('')

    # ── references ──
    refs_tex = ''
    raw_refs = content.get('references', [])
    if raw_refs:
        ref_items = []
        for idx, r in enumerate(raw_refs, 1):
            ref_items.append(r'  \bibitem{ref' + str(idx) + r'}' + '\n  ' + r)
        refs_tex = (
            r'\section*{Kaynakça / References}' + '\n'
            r'\renewcommand{\refname}{}' + '\n'
            r'\vspace{-2\baselineskip}' + '\n'
            r'\begin{thebibliography}{99}' + '\n'
            r'\setlength{\leftmargin}{1.5em}%' + '\n'
            r'\setlength{\itemindent}{-1.5em}%' + '\n\n' +
            '\n\n'.join(ref_items) + '\n\n'
            r'\end{thebibliography}'
        )
    else:
        refs_tex = (
            r'\section*{Kaynakça / References}' + '\n'
            r'\renewcommand{\refname}{}' + '\n'
            r'\vspace{-2\baselineskip}' + '\n'
            r'\begin{thebibliography}{99}' + '\n'
            r'\setlength{\leftmargin}{1.5em}%' + '\n'
            r'\setlength{\itemindent}{-1.5em}%' + '\n\n'
            r'% Kaynakları buraya ekleyin / Add your references here' + '\n\n'
            r'\end{thebibliography}'
        )

    apa_citation = (
        author_short + r'\ (' + year + r'). ' +
        tr_title + r'. \textit{Journal of Global Tourism and Technology Research}, ' +
        r'\textit{' + volume + r'}(' + issue + r'), ' +
        start_page + r'--' + end_page + r'.'
    )

    body_text = '\n'.join(body_lines) if body_lines else (
        r'\section{Giriş / Introduction}' + '\n\n'
        r'% Makale metni buraya gelecek / Article body goes here' + '\n'
    )

    doi_line = doi if doi else ''

    tex = r"""% ============================================================
%  JGTTR — Journal of Global Tourism and Technology Research
%  Bu dosya JGTTR Formatter uygulaması tarafından otomatik oluşturulmuştur.
%  Overleaf'e yükleyin: JGTTR.png ile birlikte ZIP'i açın.
%  Compile with: XeLaTeX
% ============================================================

\documentclass[10pt,a4paper]{article}

% --- Font & Unicode ---
\usepackage{fontspec}
\usepackage{unicode-math}

% --- Layout ---
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{multicol}

% --- Typography ---
\usepackage{microtype}
\usepackage{xcolor}
\usepackage{relsize}

% --- Section headings ---
\usepackage{titlesec}

% --- Tables ---
\usepackage{array}
\usepackage{tabularx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{makecell}
\usepackage{longtable}

% --- Figures ---
\usepackage{graphicx}
\usepackage{float}
\usepackage{caption}

% --- Mathematics ---
\usepackage{amsmath}

% --- References & Links ---
\usepackage{hyperref}
\usepackage{url}
\usepackage{doi}
\usepackage[numbers,sort&compress]{natbib}
\usepackage{orcidlink}

% --- Footnotes / Lists ---
\usepackage{footmisc}
\usepackage{enumitem}
\usepackage{hanging}

% --- Misc ---
\usepackage{etoolbox}
\usepackage{calc}
\usepackage{lastpage}
\usepackage{ifthen}


% ============================================================
%  FONT
% ============================================================
\setmainfont{TeX Gyre Pagella}
\setmathfont{TeX Gyre Pagella Math}


% ============================================================
%  COLOURS
% ============================================================
\definecolor{JGTTRbrown}{HTML}{833C0B}
\definecolor{JGTTRblue}{HTML}{0070C0}
\definecolor{JGTTRgray}{HTML}{CFCDCD}
\definecolor{JGTTRdarkgray}{HTML}{3B3838}


% ============================================================
%  HYPERLINKS
% ============================================================
\hypersetup{
  colorlinks=true, urlcolor=JGTTRblue,
  linkcolor=black, citecolor=black,
  pdfencoding=auto, unicode=true,
}


% ============================================================
%  GEOMETRY (body pages)
% ============================================================
\geometry{
  a4paper,
  top=1.5cm, bottom=1.5cm, left=1.5cm, right=1.5cm,
  headheight=1.2cm, headsep=0.4cm, footskip=0.8cm,
}


% ============================================================
%  PARAGRAPH FORMAT
% ============================================================
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}
\renewcommand{\baselinestretch}{1.0}


% ============================================================
%  SECTION HEADINGS
% ============================================================
\titleformat{\section}[block]{\fontsize{11}{13}\selectfont\bfseries\centering}{}{0em}{}
\titlespacing*{\section}{0pt}{9pt}{5pt}
\titleformat*{\section}{\fontsize{11}{13}\selectfont\bfseries\centering}

\titleformat{\subsection}[block]{\fontsize{11}{13}\selectfont\bfseries}{}{0em}{}
\titlespacing*{\subsection}{0pt}{7pt}{3pt}

\titleformat{\subsubsection}[block]{\fontsize{11}{13}\selectfont\bfseries\itshape}{}{0em}{}
\titlespacing*{\subsubsection}{0pt}{6pt}{3pt}


% ============================================================
%  HEADERS & FOOTERS
% ============================================================
\pagestyle{fancy}
\fancyhf{}
\fancyhead[C]{%
  \fontsize{8.5}{10.5}\selectfont\itshape
  """ + author_short + r"""\ (""" + year + r""").
  """ + head_title + r""".
  \textup{Journal of Global Tourism and Technology Research},
  \textit{""" + volume + r"""}(""" + issue + r"""),
  """ + start_page + r"""--""" + end_page + r"""%
}
\fancyfoot[C]{\fontsize{9}{11}\selectfont\thepage}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

\fancypagestyle{firstpage}{%
  \fancyhf{}%
  \renewcommand{\headrulewidth}{0pt}%
}


% ============================================================
%  CAPTION FORMAT
% ============================================================
\captionsetup{font={small,bf}, labelsep=period, justification=centering, skip=4pt}
\captionsetup[table]{position=top}


% ============================================================
%  FOOTNOTE FORMAT
% ============================================================
\renewcommand{\footnoterule}{\kern-3pt\hrule width 2cm height 0.4pt\kern 2.6pt}
\setlength{\footnotesep}{4pt}


% ============================================================
%  LIST FORMAT
% ============================================================
\setlist{leftmargin=1.2cm, itemsep=0pt, parsep=0pt, topsep=2pt, partopsep=0pt}


% ============================================================
%  ARTICLE METADATA
% ============================================================
\newcommand{\JGTTRyear}{""" + year + r"""}
\newcommand{\JGTTRvolume}{""" + volume + r"""}
\newcommand{\JGTTRissue}{""" + issue + r"""}
\newcommand{\JGTTRstartpage}{""" + start_page + r"""}
\newcommand{\JGTTRendpage}{""" + end_page + r"""}
\newcommand{\JGTTRDOI}{""" + doi_line + r"""}

\newcommand{\JGTTRarticletype}{Araştırma Makalesi -- Research Article}
\newcommand{\JGTTRreceived}{xx.xx.xxxx}
\newcommand{\JGTTRaccepted}{xx.xx.xxxx}
\newcommand{\JGTTRpublished}{xx.xx.xxxx}

\newcommand{\JGTTRturkishtitle}{""" + tr_title + r"""}
\newcommand{\JGTTRturkishabstract}{""" + escape(tr_abs) + r"""}
\newcommand{\JGTTRturkishkeywords}{""" + tr_kw + r"""}

\newcommand{\JGTTRenglishtitle}{""" + en_title + r"""}
\newcommand{\JGTTRenglishabstract}{""" + escape(en_abs) + r"""}
\newcommand{\JGTTRenglishkeywords}{""" + en_kw + r"""}

\newcommand{\JGTTRauthorshort}{""" + author_short + r"""}
\newcommand{\JGTTRheadtitle}{""" + head_title + r"""}

\newcommand{\JGTTRaffiliations}{""" + affiliations + r"""}
\newcommand{\JGTTRcorrespondinginfo}{""" + corresponding + r"""}

\newcommand{\JGTTRapacitation}{""" + apa_citation + r"""}
\newcommand{\JGTTRethicsstatement}{%
  Bu araştırma, ilgili etik kurul kararı doğrultusunda yürütülmüş olup
  tüm etik ilkelere uyulmuştur. / This study was conducted in accordance
  with the relevant ethics committee decision and all ethical principles
  were followed.%
}


% ============================================================
%  INTERNAL HELPERS
% ============================================================
\newcommand{\infolabel}[1]{{\fontsize{9}{11}\selectfont\scshape #1}}
\newcommand{\infosubheading}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\infovalue}[1]{{\fontsize{7.5}{9}\selectfont #1}}
\newcommand{\infoboldlabel}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\inforule}{%
  \par\vspace{0pt}%
  \noindent\rule{\linewidth}{0.4pt}%
  \par\vspace{0pt}%
}

% ============================================================
%  BIBLIOGRAPHY
% ============================================================
\bibliographystyle{unsrtnat}
\setlength{\bibhang}{1.5em}
\setlength{\bibsep}{3pt}


% ============================================================
%  FIRST PAGE COMMAND
% ============================================================
\newcommand{\JGTTRfirstpage}[1]{%

  \newgeometry{
    a4paper,
    top=1.2cm, bottom=1.5cm, left=1.5cm, right=1.5cm,
    headheight=0pt, headsep=0pt, footskip=0.8cm,
  }%
  \thispagestyle{firstpage}%

  % --- Logo | ISSN / URL / DOI ---
  \noindent%
  {\setlength{\tabcolsep}{0pt}%
  \setlength{\extrarowheight}{0pt}%
  \begin{tabular}{@{} m{11cm} @{} >{\raggedleft\arraybackslash}m{7cm} @{}}%
    \includegraphics[height=2.3cm]{JGTTR}%
    &%
    \raggedleft%
    {\fontsize{8.5}{11}\selectfont\color{JGTTRdarkgray}%
      \textbf{ISSN:} 2717-6924\\[2pt]%
      \href{https://dergipark.org.tr/en/pub/jgttr}{%
        \color{JGTTRblue}\itshape https://dergipark.org.tr/en/pub/jgttr%
      }\\[2pt]%
      \ifthenelse{\equal{\JGTTRDOI}{}}{%
        \textbf{DOI:} \textit{(atanacak\,/\,to be assigned)}%
      }{%
        \textbf{DOI:} \doi{\JGTTRDOI}%
      }\par%
    }%
  \end{tabular}}%

  \vspace{-1.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  {\noindent%
  \setlength{\fboxsep}{1.5pt}\setlength{\fboxrule}{0pt}%
  \colorbox{JGTTRgray}{%
    \begin{minipage}{\dimexpr\textwidth-3pt\relax}%
      {\fontsize{8}{10}\selectfont%
        \quad Year: \JGTTRyear\quad Volume: \JGTTRvolume\quad Issue: \JGTTRissue%
        \hfill%
        Yıl: \JGTTRyear\quad Cilt: \JGTTRvolume\quad Sayı: \JGTTRissue\quad%
      }%
    \end{minipage}%
  }}%

  \vspace{0.8mm}%

  {\noindent%
  \fontsize{8}{10}\selectfont\color{JGTTRdarkgray}%
  \ifthenelse{\equal{\JGTTRDOI}{}}{%
    \textbf{DOI:} \textit{(atanacak\,/\,to be assigned)}%
  }{%
    \textbf{DOI:} \doi{\JGTTRDOI}%
  }\par}%

  \vspace{1mm}%

  {\noindent\centering%
  {\fontsize{10.5}{12}\selectfont\bfseries\itshape \JGTTRarticletype}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{13}{15.5}\selectfont\bfseries \JGTTRturkishtitle}\par}%

  \vspace{0.8mm}%

  {\noindent{\fontsize{11.5}{14}\selectfont\bfseries\itshape \JGTTRenglishtitle}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{10}{12}\selectfont #1}\par}%

  \vspace{0.5mm}%

  \setlength{\tabcolsep}{4pt}%
  \noindent\rule{\textwidth}{0.4pt}%

  \noindent%
  \begin{tabular}{@{} p{0.183\textwidth} @{\hspace{2pt}} p{0.797\textwidth} @{}}%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      \centering\infolabel{MAKALE BİLGİSİ}\par%
      \inforule%
      \infosubheading{Makale Geçmişi:}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Baş. tarihi: \JGTTRreceived}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Kabul tarihi: \JGTTRaccepted}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Yayın tarihi: \JGTTRpublished}\par%
      \inforule%
      \infoboldlabel{Anahtar Kelimeler:}\par%
      \inforule%
      \infovalue{\JGTTRturkishkeywords}%
      \vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Özet}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\JGTTRturkishabstract}%
      \vspace{2pt}%
    \end{minipage}%
    \\[3pt]%

    \multicolumn{2}{@{}l@{}}{\rule{\dimexpr0.183\textwidth+0.797\textwidth+8pt\relax}{0.4pt}}\\[2pt]%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      \centering\infolabel{ARTICLE INFO}\par%
      \inforule%
      \infosubheading{Background:}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Received: \JGTTRreceived}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Accepted: \JGTTRaccepted}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Published: \JGTTRpublished}\par%
      \inforule%
      \infoboldlabel{Keywords:}\par%
      \inforule%
      \infovalue{\JGTTRenglishkeywords}%
      \vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Abstract}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\JGTTRenglishabstract}%
      \vspace{2pt}%
    \end{minipage}%
    \\[0pt]%

  \end{tabular}%

  \noindent\rule{\textwidth}{0.4pt}%

  \vspace*{\fill}%
  \noindent\begin{minipage}{\textwidth}%
  \begingroup
    \sloppy\emergencystretch=3em%
    \setlength{\parskip}{1pt}\setlength{\parindent}{0pt}%
    \fontsize{7}{8.5}\selectfont%
    \noindent\rule{\textwidth}{0.4pt}\par%
    \noindent\JGTTRaffiliations\par%
    \noindent\rule{\textwidth}{0.2pt}\par%
    \noindent\textit{*Sorumlu yazar / Corresponding author}\par%
    \noindent\textbf{Önerilen Atıf / Suggested Citation:} \JGTTRapacitation\par%
    \noindent{\setlength{\parskip}{0pt}\textbf{Etik Beyan / Ethics Statement:} \JGTTRethicsstatement\par}%
  \endgroup%
  \end{minipage}

  \restoregeometry%
}% end \JGTTRfirstpage


% ============================================================
%  BEGIN DOCUMENT
% ============================================================
\begin{document}

\JGTTRfirstpage{%
  """ + author_block + r"""%
}
% Kapak sayfası sayfa 1'dir ancak numara gösterilmez.
% İkinci sayfa sayfa 2 olarak başlar.

""" + body_text + r"""

\section*{Araştırmacıların Katkı Oranı / Author Contributions}
Kavramsal çerçeve / Conceptualization: ; Yöntem / Methodology: ;
Veri toplama / Data collection: ; Analiz / Analysis: ;
Yazım / Writing: ; Gözden geçirme / Review \& Editing:

\section*{Çıkar Çatışması / Conflict of Interest}
Yazarlar herhangi bir çıkar çatışması olmadığını beyan eder. /
The authors declare no conflict of interest.

""" + refs_tex + r"""

\end{document}
"""
    return tex


# ── ZIP builder ────────────────────────────────────────────────────────────────
def build_zip(tex_content: str, logo_src: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('main.tex', tex_content.encode('utf-8'))

        if os.path.exists(logo_src):
            zf.write(logo_src, 'JGTTR.png')

        readme = """JGTTR LaTeX Formatter — Overleaf Yükleme Rehberi
=================================================

1. Bu ZIP dosyasını açın.
2. Overleaf.com adresine gidin → New Project → Upload Project
3. ZIP dosyasının tamamını yükleyin (main.tex + JGTTR.png).
4. Derleyici olarak XeLaTeX seçin:
   Menu → Compiler → XeLaTeX
5. "Recompile" butonuna tıklayın.

Düzenleme önerileri:
- Metadata (yıl, cilt, sayı, sayfalar): dosyanın üst kısmındaki \\newcommand satırları
- Makale tarihleri: \\JGTTRreceived, \\JGTTRaccepted, \\JGTTRpublished
- Etik beyan: \\JGTTRethicsstatement
- Katkı oranları: Araştırmacıların Katkı Oranı bölümü

Sorular için: JGTTR Formatter uygulamasına bakın.
"""
        zf.writestr('README_Overleaf.txt', readme.encode('utf-8'))

        author_template = """# JGTTR Yazar Bilgileri Şablonu
# Her satır bir yazar — sütunlar: | ile ayrılır
# Sütunlar: Ad Soyad | Kurum | ORCID | E-posta | Sorumlu (evet/hayır)
# Örnek:

Ahmet Yılmaz | Turizm Bölümü, Ankara Üniversitesi, Ankara | 0000-0000-0000-0001 | ahmet@uni.edu.tr | evet
Ayşe Kaya | İşletme Bölümü, İstanbul Üniversitesi, İstanbul | 0000-0000-0000-0002 | ayse@uni.edu.tr | hayır
"""
        zf.writestr('yazar_bilgileri_sablonu.txt', author_template.encode('utf-8'))
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
#  FORM-BASED LaTeX GENERATION
# ══════════════════════════════════════════════════════════════════════════════

SECTION_CMD_MAP = {'1': r'\section', '2': r'\subsection', '3': r'\subsubsection'}

# Special starred sections (no numbering)
STARRED_NAMES = {
    'araştırmacı', 'katkı', 'contributions', 'çıkar', 'conflict',
    'teşekkür', 'acknowledgement', 'kaynakça', 'references', 'bibliography',
}


def _table_to_latex(text: str) -> str:
    """Convert pipe-separated plain text to a LaTeX tabular body."""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    if not lines:
        return ''
    rows = [[escape(c.strip()) for c in l.split('|')] for l in lines]
    ncols = max(len(r) for r in rows)
    col_spec = 'l' * ncols
    out = [r'\begin{tabular}{@{}' + col_spec + r'@{}}', r'\toprule']
    for i, row in enumerate(rows):
        # pad short rows
        while len(row) < ncols:
            row.append('')
        out.append(' & '.join(row) + r' \\')
        if i == 0:
            out.append(r'\midrule')
    out.append(r'\bottomrule')
    out.append(r'\end{tabular}')
    return '\n'.join(out)


def _build_figtable_latex(ft: dict, file_ext_map: dict) -> str:
    """Return LaTeX for one figure or table item."""
    num      = ft.get('number', '1')
    tr_cap   = escape(ft.get('tr_cap', ''))
    en_cap   = escape(ft.get('en_cap', ''))
    caption  = tr_cap + (' / ' + en_cap if en_cap else '')
    label_prefix = 'fig' if ft['type'] == 'figure' else 'tab'
    label    = label_prefix + ':' + re.sub(r'\W+', '_', num)

    if ft['type'] == 'figure':
        fkey = ft.get('file_key', '')
        fname = file_ext_map.get(fkey, 'fig_' + fkey)
        return (
            r'\begin{figure}[htbp]' + '\n'
            r'  \centering' + '\n'
            r'  \includegraphics[width=\linewidth]{' + fname + '}\n'
            r'  \caption{' + caption + '}\n'
            r'  \label{' + label + '}\n'
            r'\end{figure}'
        )
    else:
        tbl_body = _table_to_latex(ft.get('tbl_data', ''))
        if not tbl_body:
            tbl_body = '% Tablo içeriği buraya gelecek / Table content here'
        return (
            r'\begin{table}[htbp]' + '\n'
            r'  \centering' + '\n'
            r'  \caption{' + caption + '}\n'
            r'  \label{' + label + '}\n'
            '  ' + tbl_body.replace('\n', '\n  ') + '\n'
            r'\end{table}'
        )


def generate_latex_from_form(data: dict, figure_file_bytes: dict,
                             journal_settings: dict = None) -> str:
    """
    Generate complete LaTeX from structured form data.
    data keys: cover, authors, abstract, sections, figtables, extra, references
    figure_file_bytes: {file_key: (filename_in_zip, bytes)} for figure files
    journal_settings: optional journal branding/typography overrides
    """
    # ── Journal settings (with JGTTR defaults) ──
    js          = journal_settings or {}
    jname_en    = js.get('journal_name_en',  'Journal of Global Tourism and Technology Research')
    jname_tr    = js.get('journal_name_tr',  'Küresel Turizm ve Teknoloji Araştırmaları Dergisi')
    issn_print  = js.get('issn_print',  '2717-6924')
    issn_online = js.get('issn_online', '')
    j_url       = js.get('journal_url', 'https://dergipark.org.tr/en/pub/jgttr')
    font_name   = js.get('font',        'texgyrepagella')
    body_size   = js.get('body_size',   '10')
    accent_hex  = js.get('accent_color', '#833C0B').lstrip('#')
    logo_stem   = js.get('logo_stem',   'journal_logo')

    # Font setup in LaTeX
    # Kullanıcı dostu ad → (Overleaf/TeX Live display adı, math font adı veya None)
    _FONT_MAP = {
        'palatino linotype':  ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'palatino':           ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'texgyrepagella':     ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'tex gyre pagella':   ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'times new roman':    ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'times':              ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'texgyretermes':      ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'tex gyre termes':    ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'century':            ('TeX Gyre Bonum',     None),
        'century schoolbook': ('TeX Gyre Bonum',     None),
        'texgyrebonum':       ('TeX Gyre Bonum',     None),
        'calibri':            ('Carlito',            None),
        'sans serif':         ('TeX Gyre Heros',     None),
        'sans-serif':         ('TeX Gyre Heros',     None),
        'arial':              ('TeX Gyre Heros',     None),
        'texgyreheros':       ('TeX Gyre Heros',     None),
        'latinmodern':        ('Latin Modern Roman', 'Latin Modern Math'),
        'latin modern':       ('Latin Modern Roman', 'Latin Modern Math'),
    }
    _key = font_name.lower().strip()
    if _key in _FONT_MAP:
        _display, _mathfont = _FONT_MAP[_key]
        font_setup = r'\setmainfont{' + _display + '}'
        if _mathfont:
            font_setup += '\n' + r'\setmathfont{' + _mathfont + '}'
    else:
        font_setup = r'\setmainfont{' + font_name + '}'

    # ISSN display — yalnızca girilmiş alanları göster
    _issn_parts = []
    if issn_print:
        _issn_parts.append(r'\textbf{ISSN:} ' + issn_print)
    if issn_online:
        _issn_parts.append(r'\textbf{e-ISSN:} ' + issn_online)
    issn_display = (r'\ \ '.join(_issn_parts) + r'\\[2pt]%') if _issn_parts else ''

    cov      = data.get('cover', {})
    authors  = data.get('authors', [])
    abstr    = data.get('abstract', {})
    sections = data.get('sections', [])
    fts      = data.get('figtables', [])
    extra    = data.get('extra', {})
    refs_raw = data.get('references', '')

    # ── Cover fields ──
    tr_title   = escape(cov.get('tr_title', '') or 'Makalenin Türkçe Adı')
    en_title   = escape(cov.get('en_title', '') or 'Article Title in English')
    year       = escape(cov.get('year',       '2026'))
    volume     = escape(cov.get('volume',     'x'))
    issue      = escape(cov.get('issue',      'x'))
    start_page = escape(cov.get('start_page', 'xxx'))
    end_page   = escape(cov.get('end_page',   'xxx'))
    doi        = cov.get('doi', '')
    art_type   = escape(cov.get('article_type', 'Araştırma Makalesi -- Research Article'))
    received   = escape(cov.get('received',  'xx.xx.xxxx'))
    accepted   = escape(cov.get('accepted',  'xx.xx.xxxx'))
    published  = escape(cov.get('published', 'xx.xx.xxxx'))
    ethics_raw = cov.get('ethics', '').strip()
    ethics     = escape(ethics_raw) if ethics_raw else (
        r'Bu araştırma, ilgili etik kurul kararı doğrultusunda yürütülmüş olup '
        r'tüm etik ilkelere uyulmuştur. / This study was conducted in accordance '
        r'with the relevant ethics committee decision and all ethical principles were followed.'
    )

    # ── Author short / head title ──
    author_short = cov.get('author_short', '').strip()
    if not author_short and authors:
        names = [a['name'].split()[-1] for a in authors if a.get('name')]
        if len(names) == 1:
            author_short = escape(names[0])
        elif len(names) == 2:
            author_short = escape(names[0]) + r' \& ' + escape(names[1])
        elif names:
            author_short = escape(names[0]) + r' et al.'
    if not author_short:
        author_short = 'Yazar'
    else:
        author_short = escape(author_short)

    # head_title: use TR title directly (no separate field)
    head_title = tr_title

    # ── Abstract / keywords ──
    tr_abs = escape(abstr.get('tr_abs', '') or 'Türkçe özet buraya yazılmalıdır.')
    tr_kw  = escape(abstr.get('tr_kw',  '') or 'anahtar kelime 1; anahtar kelime 2')
    en_abs = escape(abstr.get('en_abs', '') or 'English abstract goes here.')
    en_kw  = escape(abstr.get('en_kw',  '') or 'keyword 1; keyword 2')

    # ── Author blocks ──
    author_block  = _format_author_block(authors)  if authors else r'Author One\textsuperscript{1,*}'
    affiliations  = _format_affiliations(authors)   if authors else r'\textsuperscript{1}Kurum Adı, Şehir.'
    corresponding = _format_corresponding(authors)  if authors else r'Yazar Adı, Kurum.'

    # ── APA citation ──
    apa_citation = (
        author_short + r'\ (' + year + r'). ' + tr_title +
        r'. \textit{' + escape(jname_en) + r'}, ' +
        r'\textit{' + volume + r'}(' + issue + r'), ' +
        start_page + r'--' + end_page + r'.'
    )

    # ── Figure file → zip name mapping ──
    # file_ext_map: {file_key: 'fig_N'} (no extension; XeLaTeX resolves)
    file_ext_map = {}
    for fkey, (zipname, _) in figure_file_bytes.items():
        # zipname like 'fig_3.png' → stem 'fig_3'
        stem = zipname.rsplit('.', 1)[0] if '.' in zipname else zipname
        file_ext_map[fkey] = stem

    # ── Body sections + paragraph-level figure/table placement ──
    body_lines = []
    placed_ft_ids = set()   # track which fts have already been placed

    for sec in sections:
        name    = sec.get('name', 'Bölüm')
        level   = sec.get('level', '1')
        content = sec.get('content', '').strip()
        cmd     = SECTION_CMD_MAP.get(level, r'\section')
        starred = any(k in name.lower() for k in STARRED_NAMES)
        star    = '*' if starred else ''
        body_lines.append(cmd + star + '{' + escape(name) + '}')
        body_lines.append('')

        # FTs assigned to this section
        sec_fts = [
            (i, ft) for i, ft in enumerate(fts)
            if ft.get('section', '').strip() == name.strip() and i not in placed_ft_ids
        ]

        if content:
            # Split content into paragraphs (blank-line separated or single newlines)
            raw_paras = [p.strip() for p in re.split(r'\n\s*\n', content)]
            raw_paras = [p for p in raw_paras if p]
            if not raw_paras:
                raw_paras = [content]

            for para in raw_paras:
                body_lines.append(escape(para))
                body_lines.append('')

                # Check if any ft's anchor text is found in this paragraph
                for i, ft in sec_fts:
                    if i in placed_ft_ids:
                        continue
                    anchor = ft.get('after_para', '').strip()
                    if anchor and anchor.lower() in para.lower():
                        body_lines.append(_build_figtable_latex(ft, file_ext_map))
                        body_lines.append('')
                        placed_ft_ids.add(i)

            # FTs for this section with anchor NOT found → append at section end
            for i, ft in sec_fts:
                if i not in placed_ft_ids:
                    body_lines.append(_build_figtable_latex(ft, file_ext_map))
                    body_lines.append('')
                    placed_ft_ids.add(i)
        else:
            # No content — place all section FTs here
            for i, ft in sec_fts:
                if i not in placed_ft_ids:
                    body_lines.append(_build_figtable_latex(ft, file_ext_map))
                    body_lines.append('')
                    placed_ft_ids.add(i)

    if not body_lines:
        body_lines = [
            r'\section{Giriş / Introduction}',
            '',
            r'% Makale metni buraya gelecek / Article body goes here',
            '',
        ]

    # ── Figures/tables with no section assigned (end of body) ──
    orphan_fts = [
        (i, ft) for i, ft in enumerate(fts)
        if not ft.get('section', '').strip() and i not in placed_ft_ids
    ]
    if orphan_fts:
        body_lines.append(r'% ── Şekil ve Tablolar / Figures and Tables ──')
        body_lines.append(r'\clearpage')
        for i, ft in orphan_fts:
            body_lines.append(_build_figtable_latex(ft, file_ext_map))
            body_lines.append('')
            placed_ft_ids.add(i)

    # ── Extra sections ──
    ack = extra.get('ack', '').strip()
    if ack:
        body_lines.append(r'\section*{Teşekkür / Acknowledgements}')
        body_lines.append(escape(ack))
        body_lines.append('')

    contrib = extra.get('contrib', '').strip()
    body_lines.append(r'\section*{Araştırmacıların Katkı Oranı / Author Contributions}')
    body_lines.append(escape(contrib) if contrib else
                      r'Kavramsal çerçeve / Conceptualization: ; Yöntem / Methodology: ; '
                      r'Veri toplama / Data collection: ; Analiz / Analysis: ; '
                      r'Yazım / Writing: ; Gözden geçirme / Review \& Editing:')
    body_lines.append('')

    conflict = extra.get('conflict', '').strip()
    body_lines.append(r'\section*{Çıkar Çatışması / Conflict of Interest}')
    body_lines.append(escape(conflict) if conflict else
                      r'Yazarlar herhangi bir çıkar çatışması olmadığını beyan eder. / '
                      r'The authors declare no conflict of interest.')
    body_lines.append('')

    # ── References — alphabetically sorted, hanging indent, no numbers ──
    refs_lines = sorted(
        [l.strip() for l in refs_raw.splitlines() if l.strip()],
        key=lambda x: x.lower()
    )
    _ref_env_open = (
        r'\section*{Kaynakça / References}' + '\n'
        r'\begin{list}{}{%' + '\n'
        r'  \setlength{\leftmargin}{1.5em}%' + '\n'
        r'  \setlength{\itemindent}{-1.5em}%' + '\n'
        r'  \setlength{\topsep}{2pt}%' + '\n'
        r'  \setlength{\itemsep}{3pt}%' + '\n'
        r'  \setlength{\parsep}{0pt}%' + '\n'
        r'}' + '\n'
    )
    if refs_lines:
        items = '\n'.join(r'\item ' + escape(r) for r in refs_lines)
        refs_tex = _ref_env_open + items + '\n' + r'\end{list}'
    else:
        refs_tex = (
            _ref_env_open +
            r'% Kaynakları buraya ekleyin / Add your references here' + '\n' +
            r'\end{list}'
        )

    body_text = '\n'.join(body_lines)
    doi_line  = doi if doi else ''

    # ── Full LaTeX document ──
    tex = r"""% ============================================================
%  """ + escape(jname_en) + r"""
%  Bu dosya Journal LaTeX Formatter tarafından oluşturulmuştur.
%  Overleaf: New Project → Upload Project → bu ZIP'i seçin.
%  Derleyici: XeLaTeX
% ============================================================

\documentclass[""" + body_size + r"""pt,a4paper]{article}

\usepackage{fontspec}
\usepackage{unicode-math}
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{microtype}
\usepackage{xcolor}
\usepackage{titlesec}
\usepackage{array}
\usepackage{tabularx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{makecell}
\usepackage{longtable}
\usepackage{graphicx}
\usepackage{float}
\usepackage{caption}
\usepackage{amsmath}
\usepackage{hyperref}
\usepackage{url}
\usepackage{doi}
\usepackage[numbers,sort&compress]{natbib}
\usepackage{orcidlink}
\usepackage{footmisc}
\usepackage{enumitem}
\usepackage{etoolbox}
\usepackage{calc}
\usepackage{lastpage}
\usepackage{ifthen}

% ── Font ──
""" + font_setup + r"""

% ── Colours ──
\definecolor{JGTTRbrown}{HTML}{""" + accent_hex + r"""}
\definecolor{JGTTRblue}{HTML}{0070C0}
\definecolor{JGTTRgray}{HTML}{CFCDCD}
\definecolor{JGTTRdarkgray}{HTML}{3B3838}

% ── Hyperlinks ──
\hypersetup{colorlinks=true,urlcolor=JGTTRblue,linkcolor=black,citecolor=black,pdfencoding=auto,unicode=true}
% URL'leri bölme — sığmazsa bütün olarak alt satıra geç
\renewcommand{\UrlBreaks}{}
\renewcommand{\UrlBigBreaks}{}

% ── Geometry ──
\geometry{a4paper,top=1.5cm,bottom=1.5cm,left=1.5cm,right=1.5cm,headheight=1.2cm,headsep=0.4cm,footskip=0.8cm}

% ── Paragraph format ──
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}
\renewcommand{\baselinestretch}{1.0}

% ── Satır kırma / Line breaking ──
% Uzun kelimeler ve URL'lerin sayfa kenarına taşmasını önler
\setlength{\emergencystretch}{3em}
\tolerance=800
\hyphenpenalty=50
\exhyphenpenalty=50

% ── Section headings ──
\titleformat{\section}[block]{\fontsize{11}{13}\selectfont\bfseries\centering}{}{0em}{}
\titlespacing*{\section}{0pt}{9pt}{5pt}
\titleformat*{\section}{\fontsize{11}{13}\selectfont\bfseries\centering}
\titleformat{\subsection}[block]{\fontsize{11}{13}\selectfont\bfseries}{}{0em}{}
\titlespacing*{\subsection}{0pt}{7pt}{3pt}
\titleformat{\subsubsection}[block]{\fontsize{11}{13}\selectfont\bfseries\itshape}{}{0em}{}
\titlespacing*{\subsubsection}{0pt}{6pt}{3pt}

% ── Headers & footers ──
\pagestyle{fancy}
\fancyhf{}
\fancyhead[C]{%
  \fontsize{8.5}{10.5}\selectfont
  """ + author_short + r"""\ (""" + year + r""").
  """ + head_title + r""".
  \textit{""" + escape(jname_en) + r"""},
  \textit{""" + volume + r"""}(""" + issue + r"""),
  """ + start_page + r"""--""" + end_page + r"""%
}
\fancyfoot[C]{\fontsize{9}{11}\selectfont\thepage}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}
\fancypagestyle{firstpage}{\fancyhf{}\renewcommand{\headrulewidth}{0pt}}

% ── Caption format ──
\captionsetup{font={small,bf},labelsep=period,justification=centering,skip=4pt}
\captionsetup[table]{position=top}

% ── Footnote format ──
\renewcommand{\footnoterule}{\kern-3pt\hrule width 2cm height 0.4pt\kern 2.6pt}
\setlength{\footnotesep}{4pt}

% ── List format ──
\setlist{leftmargin=1.2cm,itemsep=0pt,parsep=0pt,topsep=2pt,partopsep=0pt}

% ── Bibliography ──
\bibliographystyle{unsrtnat}
\setlength{\bibhang}{1.5em}
\setlength{\bibsep}{3pt}

% ── Article metadata ──
\newcommand{\JGTTRyear}{""" + year + r"""}
\newcommand{\JGTTRvolume}{""" + volume + r"""}
\newcommand{\JGTTRissue}{""" + issue + r"""}
\newcommand{\JGTTRstartpage}{""" + start_page + r"""}
\newcommand{\JGTTRendpage}{""" + end_page + r"""}
\newcommand{\JGTTRDOI}{""" + doi_line + r"""}
\newcommand{\JGTTRarticletype}{""" + art_type + r"""}
\newcommand{\JGTTRreceived}{""" + received + r"""}
\newcommand{\JGTTRaccepted}{""" + accepted + r"""}
\newcommand{\JGTTRpublished}{""" + published + r"""}
\newcommand{\JGTTRturkishtitle}{""" + tr_title + r"""}
\newcommand{\JGTTRturkishabstract}{""" + tr_abs + r"""}
\newcommand{\JGTTRturkishkeywords}{""" + tr_kw + r"""}
\newcommand{\JGTTRenglishtitle}{""" + en_title + r"""}
\newcommand{\JGTTRenglishabstract}{""" + en_abs + r"""}
\newcommand{\JGTTRenglishkeywords}{""" + en_kw + r"""}
\newcommand{\JGTTRauthorshort}{""" + author_short + r"""}
\newcommand{\JGTTRheadtitle}{""" + head_title + r"""}
\newcommand{\JGTTRaffiliations}{""" + affiliations + r"""}
\newcommand{\JGTTRcorrespondinginfo}{""" + corresponding + r"""}
\newcommand{\JGTTRapacitation}{""" + apa_citation + r"""}
\newcommand{\JGTTRethicsstatement}{""" + ethics + r"""}

% ── Internal helpers ──
\newcommand{\infolabel}[1]{{\fontsize{9}{11}\selectfont\scshape #1}}
\newcommand{\infosubheading}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\infovalue}[1]{{\fontsize{7.5}{9}\selectfont #1}}
\newcommand{\infoboldlabel}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\inforule}{\par\vspace{0pt}\noindent\rule{\linewidth}{0.4pt}\par\vspace{0pt}}

% ── First page command ──
\newcommand{\JGTTRfirstpage}[1]{%
  \newgeometry{a4paper,top=1.2cm,bottom=1.5cm,left=1.5cm,right=1.5cm,headheight=0pt,headsep=0pt,footskip=0.8cm}%
  \thispagestyle{firstpage}%

  \noindent%
  {\setlength{\tabcolsep}{0pt}\setlength{\extrarowheight}{0pt}%
  \begin{tabular}{@{} m{11cm} @{} >{\raggedleft\arraybackslash}m{7cm} @{}}%
    \includegraphics[height=2.3cm]{""" + logo_stem + r"""}%
    &%
    \raggedleft%
    {\fontsize{8.5}{11}\selectfont\color{JGTTRdarkgray}%
      """ + issn_display + r"""
      \href{""" + j_url + r"""}{\color{JGTTRblue}\itshape """ + j_url + r"""}\\[2pt]%
      \ifthenelse{\equal{\JGTTRDOI}{}}{\textbf{DOI:} \textit{(atanacak\,/\,to be assigned)}}{\textbf{DOI:} \doi{\JGTTRDOI}}\par}%
  \end{tabular}}%

  \vspace{-1.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  {\noindent\setlength{\fboxsep}{1.5pt}\setlength{\fboxrule}{0pt}%
  \colorbox{JGTTRgray}{%
    \begin{minipage}{\dimexpr\textwidth-3pt\relax}%
      {\fontsize{8}{10}\selectfont\quad Year: \JGTTRyear\quad Volume: \JGTTRvolume\quad Issue: \JGTTRissue%
       \hfill Yıl: \JGTTRyear\quad Cilt: \JGTTRvolume\quad Sayı: \JGTTRissue\quad}%
    \end{minipage}%
  }}%

  \vspace{1.5mm}%

  {\noindent\centering{\fontsize{10.5}{12}\selectfont\bfseries\itshape \JGTTRarticletype}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{13}{15.5}\selectfont\bfseries \JGTTRturkishtitle}\par}%

  \vspace{0.8mm}%

  {\noindent{\fontsize{11.5}{14}\selectfont\bfseries\itshape \JGTTRenglishtitle}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{10}{12}\selectfont #1}\par}%

  \vspace{0.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  \noindent%
  \begin{tabular}{@{} p{0.183\textwidth} @{\hspace{2pt}} p{0.797\textwidth} @{}}%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%
      \centering\infolabel{MAKALE BİLGİSİ}\par\inforule%
      \infosubheading{Makale Geçmişi:}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Baş. tarihi: \JGTTRreceived}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Kabul tarihi: \JGTTRaccepted}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Yayın tarihi: \JGTTRpublished}\par\inforule%
      \infoboldlabel{Anahtar Kelimeler:}\par\inforule%
      \infovalue{\JGTTRturkishkeywords}\vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Özet}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\JGTTRturkishabstract}\vspace{2pt}%
    \end{minipage}%
    \\[3pt]%

    \multicolumn{2}{@{}l@{}}{\rule{\dimexpr0.183\textwidth+0.797\textwidth+8pt\relax}{0.4pt}}\\[2pt]%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%
      \centering\infolabel{ARTICLE INFO}\par\inforule%
      \infosubheading{Background:}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Received: \JGTTRreceived}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Accepted: \JGTTRaccepted}\par\inforule%
      {\fontsize{7.5}{9}\selectfont Published: \JGTTRpublished}\par\inforule%
      \infoboldlabel{Keywords:}\par\inforule%
      \infovalue{\JGTTRenglishkeywords}\vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Abstract}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\JGTTRenglishabstract}\vspace{2pt}%
    \end{minipage}%
    \\[0pt]%

  \end{tabular}%

  \noindent\rule{\textwidth}{0.4pt}%

  \vspace*{\fill}%
  \noindent\begin{minipage}{\textwidth}%
  \begingroup
    \sloppy\emergencystretch=3em%
    \setlength{\parskip}{1pt}\setlength{\parindent}{0pt}%
    \fontsize{7}{8.5}\selectfont%
    \noindent\rule{\textwidth}{0.4pt}\par%
    \noindent\JGTTRaffiliations\par%
    \noindent\rule{\textwidth}{0.2pt}\par%
    \noindent\textit{*Sorumlu yazar / Corresponding author}\par%
    \noindent\textbf{Önerilen Atıf / Suggested Citation:} \JGTTRapacitation\par%
    \noindent{\setlength{\parskip}{0pt}\textbf{Etik Beyan / Ethics Statement:} \JGTTRethicsstatement\par}%
  \endgroup%
  \end{minipage}

  \restoregeometry%
}% end \JGTTRfirstpage


% ============================================================
\begin{document}

\JGTTRfirstpage{%
  """ + author_block + r"""%
}
% Kapak sayfası sayfa 1'dir ancak numara gösterilmez.
% İkinci sayfa sayfa 2 olarak başlar.

""" + body_text + '\n\n' + refs_tex + r"""

\end{document}
"""
    return tex


def build_zip_form(tex_content: str, logo_src: str, figure_file_bytes: dict,
                   journal_settings: dict = None) -> bytes:
    """
    Build Overleaf-ready ZIP including:
    - main.tex
    - journal_logo.<ext> (logo)
    - figure files (fig_N.ext)
    - README_Overleaf.txt
    """
    js = journal_settings or {}
    logo_stem = js.get('logo_stem', 'journal_logo')
    logo_fn   = js.get('logo_filename', '')   # e.g. "JGTTR_logo.png"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('main.tex', tex_content.encode('utf-8'))

        if os.path.exists(logo_src):
            # Determine extension of the actual logo file
            ext = logo_src.rsplit('.', 1)[-1].lower() if '.' in logo_src else 'png'
            zf.write(logo_src, logo_stem + '.' + ext)

        for fkey, (zipname, filebytes) in figure_file_bytes.items():
            zf.writestr(zipname, filebytes)

        readme = (
            "JGTTR LaTeX Formatter — Overleaf Yükleme Rehberi\n"
            "=================================================\n\n"
            "1. Bu ZIP dosyasını açın.\n"
            "2. Overleaf.com → New Project → Upload Project → ZIP'i seçin.\n"
            "3. Menu → Compiler → XeLaTeX seçin.\n"
            "4. Recompile → PDF hazır.\n"
        )
        zf.writestr('README_Overleaf.txt', readme.encode('utf-8'))

    buf.seek(0)
    return buf.read()
